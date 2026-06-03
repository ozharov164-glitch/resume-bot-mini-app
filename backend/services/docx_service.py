"""DOCX export mirroring PDF resume templates (classic / modern / compact)."""

from __future__ import annotations

import io
import re
from typing import Any

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from services.docx_font_embed import embed_nunito_fonts
from services.resume_schema import normalize_resume_data
from services.resume_text_utils import split_bullets

FONT_NAME = "Nunito Sans"
VALID_TEMPLATES = frozenset({"classic", "modern", "compact"})

# Classic PDF palette — pixel-matched to get_pdf_styles() / resume_classic.html
C_SIDEBAR_BG = "0D1F14"
C_SIDEBAR_TITLE = "2DE08A"
C_SIDEBAR_TEXT = "D1D5DB"       # rgba(255,255,255,0.82)
C_SIDEBAR_MUTED = "999999"      # rgba(255,255,255,0.40)
C_CHIP_BG = "12241C"            # rgba(45,224,138,0.09) on dark
C_CHIP_BORDER = "2A4A38"        # rgba(45,224,138,0.18)
C_SALARY_BG = "142118"          # rgba(45,224,138,0.10) on dark
C_SALARY_BORDER = "2A4536"      # rgba(45,224,138,0.22)
C_TEXT_DARK = "0D1F14"
C_TEXT_BODY = "374151"
C_TEXT_MUTED = "6B7280"
C_BRAND_GREEN = "16A34A"        # main-position in PDF
C_SUMMARY_BG = "F7FDF9"
C_ACCENT = "2DE08A"

# Modern PDF palette
M_ACCENT = "2563EB"
M_CHIP_BG = "EFF6FF"
M_CHIP_BORDER = "BFDBFE"
M_CHIP_TEXT = "1E40AF"

# Compact PDF palette
K_ACCENT = "7C3AED"
K_SIDEBAR_BG = "F8F8F8"
K_SALARY_BG = "F3E8FF"


def _rgb(hex_color: str) -> RGBColor:
    h = hex_color.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _normalize_salary(resume_data: dict[str, Any]) -> dict[str, Any]:
    data = dict(resume_data)
    salary = data.get("salary", "")
    if salary:
        salary_clean = re.sub(r"[^\d\s]", "", str(salary)).strip()
        if salary_clean:
            data["salary"] = salary_clean + " ₽/мес"
    return data


def _set_cell_bg(cell, fill_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex.upper())
    tc_pr.append(shd)


def _set_cell_margins(cell, *, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = OxmlElement("w:tcMar")
    for side, val in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        tc_mar.append(node)
    tc_pr.append(tc_mar)


def _remove_table_borders(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    if tbl.tblPr is None:
        tbl.insert(0, tbl_pr)
    borders = OxmlElement("w:tblBorders")
    for name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge = OxmlElement(f"w:{name}")
        edge.set(qn("w:val"), "nil")
        borders.append(edge)
    tbl_pr.append(borders)


def _shade_paragraph(paragraph, fill_hex: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex.upper())
    p_pr.append(shd)


def _paragraph_border(
    paragraph,
    *,
    left: str | None = None,
    bottom: str | None = None,
    size: int = 8,
    color: str = C_ACCENT,
) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    if left:
        side = OxmlElement("w:left")
        side.set(qn("w:val"), "single")
        side.set(qn("w:sz"), str(size * 2))
        side.set(qn("w:color"), color.upper())
        side.set(qn("w:space"), "4")
        p_bdr.append(side)
    if bottom:
        side = OxmlElement("w:bottom")
        side.set(qn("w:val"), "single")
        side.set(qn("w:sz"), str(size))
        side.set(qn("w:color"), color.upper())
        side.set(qn("w:space"), "2")
        p_bdr.append(side)
    p_pr.append(p_bdr)


def _style_run(
    run,
    *,
    size_pt: float,
    bold: bool = False,
    italic: bool = False,
    color: str | None = None,
    font_name: str = FONT_NAME,
) -> None:
    run.bold = bold
    run.italic = italic
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    if color:
        run.font.color.rgb = _rgb(color)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
        r_fonts.set(qn(f"w:{attr}"), font_name)


def _hh_contact_line(data: dict[str, Any]) -> str:
    """Single-line contacts for hh.ru / ATS parsers (visible in main column)."""
    parts: list[str] = []
    city = str(data.get("city") or "").strip()
    phone = str(data.get("phone") or "").strip()
    email = str(data.get("email") or "").strip()
    if city:
        parts.append(city)
    if phone:
        parts.append(phone)
    if email:
        parts.append(email)
    return " · ".join(parts)


def _apply_doc_properties(doc: Document, data: dict[str, Any]) -> None:
    name = str(data.get("full_name") or "").strip()
    position = str(data.get("target_position") or "").strip()
    if name:
        doc.core_properties.title = name
    if position:
        doc.core_properties.subject = position
    doc.core_properties.category = "hh.ru"
    doc.core_properties.keywords = "резюме, hh.ru, ResumeBot"


def docx_filename(resume_data: dict[str, Any], template_name: str = "classic") -> str:
    """hh.ru-friendly attachment name with template slug."""
    name = str(resume_data.get("full_name") or "resume").strip()
    safe = re.sub(r"[^\w\s-]", "", name, flags=re.UNICODE)
    safe = re.sub(r"\s+", "_", safe).strip("_")[:60] or "resume"
    tpl = (template_name or "classic").strip().lower()
    if tpl not in VALID_TEMPLATES:
        tpl = "classic"
    return f"Rezyume_{safe}_{tpl}_hh.docx"


def _set_cell_borders(
    cell,
    *,
    top: str | None = None,
    bottom: str | None = None,
    left: str | None = None,
    right: str | None = None,
    color: str = C_ACCENT,
    size: int = 4,
) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for side_name, val in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        if val:
            edge = OxmlElement(f"w:{side_name}")
            edge.set(qn("w:val"), val)
            edge.set(qn("w:sz"), str(size))
            edge.set(qn("w:color"), color.upper())
            edge.set(qn("w:space"), "0")
            borders.append(edge)
    tc_pr.append(borders)


def _set_table_fixed_width(table, *, pct: int = 5000) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    if tbl.tblPr is None:
        tbl.insert(0, tbl_pr)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)
    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:w"), str(pct))
    tbl_w.set(qn("w:type"), "pct")
    tbl_pr.append(tbl_w)


def _set_row_min_height(row, twips: int = 15840) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tr_height = OxmlElement("w:trHeight")
    tr_height.set(qn("w:val"), str(twips))
    tr_height.set(qn("w:hRule"), "atLeast")
    tr_pr.append(tr_height)


def _clear_paragraph_spacing(paragraph) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing = 1.15


def _para_divider(container, *, color: str = C_TEXT_DARK, size: int = 16, space_after: float = 14) -> None:
    p = container.add_paragraph()
    _clear_paragraph_spacing(p)
    p.paragraph_format.space_after = Pt(space_after)
    _paragraph_border(p, bottom=color, size=size, color=color)


def _add_sidebar_title(
    cell,
    text: str,
    *,
    title_color: str = C_SIDEBAR_TITLE,
    border_color: str = "2A4A38",
) -> None:
    p = cell.add_paragraph()
    _clear_paragraph_spacing(p)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(text.upper())
    _style_run(run, size_pt=6.5, bold=True, color=title_color)
    _paragraph_border(p, bottom=border_color, size=4, color=border_color)


def _add_light_sidebar_title(cell, text: str, *, accent: str = K_ACCENT) -> None:
    p = cell.add_paragraph()
    _clear_paragraph_spacing(p)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text.upper())
    _style_run(run, size_pt=6, bold=True, color=accent)
    _paragraph_border(p, bottom="E5E7EB", size=4, color="E5E7EB")


def _add_main_title(cell, text: str, *, accent: str = C_ACCENT, dark: str = C_TEXT_DARK) -> None:
    """Section title with short green underline (inline-block effect via nested table)."""
    outer = cell.add_table(rows=1, cols=1)
    _remove_table_borders(outer)
    _set_table_fixed_width(outer, pct=1800)
    sc = outer.rows[0].cells[0]
    _set_cell_margins(sc, top=0, start=0, bottom=60, end=0)
    _set_cell_borders(sc, bottom="single", color=accent, size=10)
    p = sc.paragraphs[0]
    _clear_paragraph_spacing(p)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text.upper())
    _style_run(run, size_pt=7.5, bold=True, color=dark)
    gap = cell.add_paragraph()
    gap.paragraph_format.space_after = Pt(6)


def _add_sidebar_label_value(cell, label: str, value: str, *, muted: str = C_SIDEBAR_MUTED, text: str = C_SIDEBAR_TEXT) -> None:
    p = cell.add_paragraph()
    _clear_paragraph_spacing(p)
    p.paragraph_format.space_after = Pt(4)
    label_run = p.add_run(label.upper() + "\n")
    _style_run(label_run, size_pt=6.5, bold=True, color=muted)
    value_run = p.add_run(value)
    _style_run(value_run, size_pt=8, color=text)


def _add_sidebar_dot_line(cell, text: str) -> None:
    p = cell.add_paragraph()
    _clear_paragraph_spacing(p)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.25)
    p.paragraph_format.first_line_indent = Cm(-0.25)
    dot = p.add_run("· ")
    _style_run(dot, size_pt=9, bold=True, color=C_SIDEBAR_TITLE)
    body = p.add_run(text)
    _style_run(body, size_pt=8, color=C_SIDEBAR_TEXT)


def _add_cert_line(cell, text: str) -> None:
    p = cell.add_paragraph()
    _clear_paragraph_spacing(p)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.35)
    p.paragraph_format.first_line_indent = Cm(-0.35)
    mark = p.add_run("✓ ")
    _style_run(mark, size_pt=7, bold=True, color=C_SIDEBAR_TITLE)
    body = p.add_run(text)
    _style_run(body, size_pt=7.5, color="B8CFC0")


def _add_salary_box(cell, salary: str) -> None:
    tbl = cell.add_table(rows=1, cols=1)
    _remove_table_borders(tbl)
    _set_table_fixed_width(tbl, pct=4800)
    box = tbl.rows[0].cells[0]
    _set_cell_bg(box, C_SALARY_BG)
    _set_cell_borders(box, top="single", bottom="single", left="single", right="single", color=C_SALARY_BORDER, size=4)
    _set_cell_margins(box, top=80, start=100, bottom=80, end=100)
    p = box.paragraphs[0]
    _clear_paragraph_spacing(p)
    val = p.add_run(salary + "\n")
    _style_run(val, size_pt=11.5, bold=True, color=C_SIDEBAR_TITLE)
    lbl = p.add_run("Желаемая зарплата")
    _style_run(lbl, size_pt=6.5, color=C_SIDEBAR_MUTED)
    spacer = cell.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)


def _add_section_title_row(
    container,
    text: str,
    *,
    accent: str,
    dark: str = C_TEXT_DARK,
    font_pt: float = 7.5,
    border_sz: int = 12,
    space_before: float = 10,
) -> None:
    """Full-width section title (modern template)."""
    p = container.add_paragraph()
    _clear_paragraph_spacing(p)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(6)
    _style_run(p.add_run(text.upper()), size_pt=font_pt, bold=True, color=dark)
    _paragraph_border(p, bottom=accent, size=border_sz, color=accent)


def _add_skills_chips(
    cell,
    skills: list[str],
    *,
    bg: str = C_CHIP_BG,
    border: str = C_CHIP_BORDER,
    text: str = C_SIDEBAR_TEXT,
    font_pt: float = 7,
    cols: int = 2,
) -> None:
    if not skills:
        return
    rows_needed = (len(skills) + cols - 1) // cols
    tbl = cell.add_table(rows=rows_needed, cols=cols)
    _remove_table_borders(tbl)
    _set_table_fixed_width(tbl, pct=4800)
    idx = 0
    for r in range(rows_needed):
        for c in range(cols):
            chip_cell = tbl.rows[r].cells[c]
            _set_cell_margins(chip_cell, top=20, start=20, bottom=20, end=20)
            if idx >= len(skills):
                continue
            skill = skills[idx]
            idx += 1
            _set_cell_bg(chip_cell, bg)
            _set_cell_borders(chip_cell, top="single", bottom="single", left="single", right="single", color=border, size=2)
            p = chip_cell.paragraphs[0]
            _clear_paragraph_spacing(p)
            _style_run(p.add_run(skill), size_pt=font_pt, color=text)
    gap = cell.add_paragraph()
    gap.paragraph_format.space_after = Pt(2)


def _add_compact_salary_box(cell, salary: str) -> None:
    tbl = cell.add_table(rows=1, cols=1)
    _remove_table_borders(tbl)
    _set_table_fixed_width(tbl, pct=4800)
    box = tbl.rows[0].cells[0]
    _set_cell_bg(box, K_SALARY_BG)
    _set_cell_borders(box, left="single", color=K_ACCENT, size=16)
    _set_cell_margins(box, top=70, start=100, bottom=70, end=90)
    p = box.paragraphs[0]
    _clear_paragraph_spacing(p)
    _style_run(p.add_run(salary + "\n"), size_pt=10, bold=True, color=K_ACCENT)
    _style_run(p.add_run("Желаемая зарплата"), size_pt=6, color="6B7280")


def _add_experience_entries(
    cell,
    experience: list[Any],
    *,
    accent: str,
    position_color: str,
    company_color: str = C_TEXT_DARK,
    tab_cm: float = 11.0,
    company_pt: float = 9.5,
    position_pt: float = 8.5,
    bullet_pt: float = 8.5,
) -> None:
    for idx, job in enumerate(experience):
        if not isinstance(job, dict):
            continue
        company = str(job.get("company") or "").strip()
        period = str(job.get("period") or "").strip()
        position_j = str(job.get("position") or "").strip()
        if company or period:
            p = cell.add_paragraph()
            _clear_paragraph_spacing(p)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.tab_stops.add_tab_stop(Cm(tab_cm), WD_TAB_ALIGNMENT.RIGHT)
            if company:
                _style_run(p.add_run(company), size_pt=company_pt, bold=True, color=company_color)
            if period:
                p.add_run("\t")
                _style_run(p.add_run(period), size_pt=7, italic=True, color=C_TEXT_MUTED)
        if position_j:
            p_pos = cell.add_paragraph()
            _clear_paragraph_spacing(p_pos)
            p_pos.paragraph_format.space_after = Pt(4)
            _style_run(p_pos.add_run(position_j), size_pt=position_pt, bold=True, color=position_color)
        for bullet in split_bullets(str(job.get("description") or "")):
            if bullet:
                _add_bullet_line(cell, bullet, marker=accent, size_pt=bullet_pt)
        if idx < len(experience) - 1:
            sep = cell.add_paragraph()
            _clear_paragraph_spacing(sep)
            sep.paragraph_format.space_before = Pt(4)
            sep.paragraph_format.space_after = Pt(8)
            _paragraph_border(sep, bottom="F0F0F0", size=4, color="F0F0F0")


def _add_education_entries(cell, education: list[Any], *, institution_pt: float = 9, details_pt: float = 8) -> None:
    for edu in education:
        if not isinstance(edu, dict):
            continue
        institution = str(edu.get("institution") or "").strip()
        degree = str(edu.get("degree") or "").strip()
        year = str(edu.get("year") or "").strip()
        if not institution and not degree:
            continue
        if institution:
            p = cell.add_paragraph()
            _clear_paragraph_spacing(p)
            _style_run(p.add_run(institution), size_pt=institution_pt, bold=True, color=C_TEXT_DARK)
        details = degree
        if year and year != "0":
            details = f"{details} · {year}" if details else year
        if details:
            p2 = cell.add_paragraph()
            _clear_paragraph_spacing(p2)
            p2.paragraph_format.space_after = Pt(6)
            _style_run(p2.add_run(details), size_pt=details_pt, color=C_TEXT_MUTED)


def _add_achievement_entries(cell, items: list[str], *, accent: str, size_pt: float = 8.5) -> None:
    for item in items:
        _add_bullet_line(cell, item, marker=accent, size_pt=size_pt)


def _add_bullet_line(cell, text: str, *, color: str = C_TEXT_BODY, marker: str = C_ACCENT, size_pt: float = 8.5) -> None:
    p = cell.add_paragraph()
    _clear_paragraph_spacing(p)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.4)
    p.paragraph_format.first_line_indent = Cm(-0.4)
    dot = p.add_run("· ")
    _style_run(dot, size_pt=11, bold=True, color=marker)
    body = p.add_run(text)
    _style_run(body, size_pt=size_pt, color=color)


def _add_spacer(cell, pt: float = 6) -> None:
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(pt)


def _cell_all_text(cell) -> str:
    parts = [cell.text]
    for tbl in cell.tables:
        for row in tbl.rows:
            for c in row.cells:
                parts.append(_cell_all_text(c))
    return "\n".join(parts)


def _schedule_text(resume_data: dict[str, Any]) -> str:
    schedule = resume_data.get("work_schedule")
    if not schedule:
        return ""
    if isinstance(schedule, list):
        return ", ".join(str(s).strip() for s in schedule if str(s).strip())
    return str(schedule).strip()
    schedule = resume_data.get("work_schedule")
    if not schedule:
        return ""
    if isinstance(schedule, list):
        return ", ".join(str(s).strip() for s in schedule if str(s).strip())
    return str(schedule).strip()


def _non_native_languages(resume_data: dict[str, Any]) -> list[str]:
    langs = resume_data.get("languages") or []
    return [str(lang).strip() for lang in langs if str(lang).strip() and str(lang).strip() != "Русский — родной"]


def _fill_classic_sidebar(cell, data: dict[str, Any]) -> None:
    _set_cell_bg(cell, C_SIDEBAR_BG)
    _set_cell_margins(cell, top=180, start=180, bottom=180, end=160)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

    _add_sidebar_title(cell, "Контакты")
    if data.get("phone"):
        _add_sidebar_label_value(cell, "Телефон", str(data["phone"]))
    if data.get("email"):
        _add_sidebar_label_value(cell, "Email", str(data["email"]))
    if data.get("city"):
        _add_sidebar_label_value(cell, "Город", str(data["city"]))

    salary = str(data.get("salary") or "").strip()
    if salary:
        _add_salary_box(cell, salary)

    sched = _schedule_text(data)
    if sched:
        _add_sidebar_title(cell, "График")
        _add_sidebar_dot_line(cell, sched)

    relocation = str(data.get("relocation") or "").strip()
    if relocation:
        _add_sidebar_title(cell, "Переезд")
        _add_sidebar_dot_line(cell, relocation)

    skills = [str(s).strip() for s in (data.get("skills") or []) if str(s).strip()]
    if skills:
        _add_sidebar_title(cell, "Навыки")
        _add_skills_chips(cell, skills)

    langs = _non_native_languages(data)
    if langs:
        _add_sidebar_title(cell, "Языки")
        for lang in data.get("languages") or []:
            lang_s = str(lang).strip()
            if lang_s:
                _add_sidebar_dot_line(cell, lang_s)

    docs = [str(d).strip() for d in (data.get("documents_and_permits") or []) if str(d).strip()]
    if docs:
        _add_sidebar_title(cell, "Документы и допуски")
        for doc_line in docs:
            _add_cert_line(cell, doc_line)


def _fill_classic_main(cell, data: dict[str, Any]) -> None:
    _set_cell_margins(cell, top=180, start=160, bottom=180, end=180)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

    name = str(data.get("full_name") or "").strip()
    position = str(data.get("target_position") or "").strip()
    if name:
        p = cell.add_paragraph()
        _clear_paragraph_spacing(p)
        p.paragraph_format.space_after = Pt(2)
        _style_run(p.add_run(name), size_pt=19, bold=True, color=C_TEXT_DARK)
    if position:
        p = cell.add_paragraph()
        _clear_paragraph_spacing(p)
        p.paragraph_format.space_after = Pt(3)
        _style_run(p.add_run(position), size_pt=10, bold=True, color=C_BRAND_GREEN)
    contact = _hh_contact_line(data)
    if contact:
        p = cell.add_paragraph()
        _clear_paragraph_spacing(p)
        p.paragraph_format.space_after = Pt(6)
        _style_run(p.add_run(contact), size_pt=8, color=C_TEXT_MUTED)
    _para_divider(cell, color=C_TEXT_DARK, size=16, space_after=14)

    summary = str(data.get("summary") or "").strip()
    if summary:
        _add_main_title(cell, "О себе")
        tbl = cell.add_table(rows=1, cols=1)
        _remove_table_borders(tbl)
        _set_table_fixed_width(tbl, pct=4800)
        box = tbl.rows[0].cells[0]
        _set_cell_bg(box, C_SUMMARY_BG)
        _set_cell_borders(box, left="single", color=C_ACCENT, size=12)
        _set_cell_margins(box, top=100, start=120, bottom=100, end=120)
        p = box.paragraphs[0]
        _clear_paragraph_spacing(p)
        _style_run(p.add_run(summary), size_pt=9, color=C_TEXT_BODY)
        gap = cell.add_paragraph()
        gap.paragraph_format.space_after = Pt(8)

    experience = data.get("experience") or []
    if experience:
        _add_main_title(cell, "Опыт работы")
        _add_experience_entries(
            cell,
            experience,
            accent=C_ACCENT,
            position_color=C_BRAND_GREEN,
        )

    achievements = [str(a).strip() for a in (data.get("key_achievements") or []) if str(a).strip()]
    if achievements:
        _add_main_title(cell, "Ключевые достижения")
        _add_achievement_entries(cell, achievements, accent=C_ACCENT)

    education = data.get("education") or []
    if education:
        _add_main_title(cell, "Образование")
        _add_education_entries(cell, education)


def _apply_default_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = FONT_NAME
    normal.font.size = Pt(9)
    pf = normal.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.15


def _build_classic_docx(data: dict[str, Any]) -> bytes:
    doc = Document()
    _apply_default_styles(doc)
    _apply_doc_properties(doc, data)
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(0.8)
    section.right_margin = Cm(0.8)
    section.top_margin = Cm(0.8)
    section.bottom_margin = Cm(0.8)

    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    _remove_table_borders(table)
    _set_table_fixed_width(table)
    _set_row_min_height(table.rows[0])
    sidebar, main = table.rows[0].cells
    sidebar.width = Cm(5.9)
    main.width = Cm(13.1)

    _fill_classic_sidebar(sidebar, data)
    _fill_classic_main(main, data)
    return _save(doc)


def _build_modern_docx_clean(data: dict[str, Any]) -> bytes:
    doc = Document()
    _apply_default_styles(doc)
    _apply_doc_properties(doc, data)
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(1.6)
    section.right_margin = Cm(1.6)
    section.top_margin = Cm(1.4)
    section.bottom_margin = Cm(1.8)

    name = str(data.get("full_name") or "").strip()
    if name:
        p = doc.add_paragraph()
        _clear_paragraph_spacing(p)
        p.paragraph_format.space_after = Pt(4)
        _style_run(p.add_run(name), size_pt=24, bold=True, color="111827")

    meta_parts: list[str] = []
    for key in ("target_position", "city", "phone", "email"):
        val = str(data.get(key) or "").strip()
        if val:
            meta_parts.append(val)
    if meta_parts:
        p = doc.add_paragraph()
        _clear_paragraph_spacing(p)
        p.paragraph_format.space_after = Pt(4)
        _style_run(p.add_run(" · ".join(meta_parts)), size_pt=8.5, color="4B5563")

    salary = str(data.get("salary") or "").strip()
    if salary:
        p = doc.add_paragraph()
        _clear_paragraph_spacing(p)
        p.paragraph_format.space_after = Pt(3)
        _style_run(p.add_run("Зарплата: "), size_pt=8.5, color="4B5563")
        _style_run(p.add_run(salary), size_pt=8.5, bold=True, color="111827")

    sched = _schedule_text(data)
    if sched:
        p = doc.add_paragraph()
        _clear_paragraph_spacing(p)
        p.paragraph_format.space_after = Pt(3)
        _style_run(p.add_run("График: "), size_pt=8.5, color="4B5563")
        _style_run(p.add_run(sched), size_pt=8.5, bold=True, color="111827")

    relocation = str(data.get("relocation") or "").strip()
    if relocation:
        p = doc.add_paragraph()
        _clear_paragraph_spacing(p)
        p.paragraph_format.space_after = Pt(6)
        _style_run(p.add_run("Переезд: "), size_pt=8.5, color="4B5563")
        _style_run(p.add_run(relocation), size_pt=8.5, bold=True, color="111827")

    _para_divider(doc, color=M_ACCENT, size=16, space_after=14)

    summary = str(data.get("summary") or "").strip()
    if summary:
        _add_section_title_row(doc, "О себе", accent=M_ACCENT, dark="111827")
        p = doc.add_paragraph()
        _clear_paragraph_spacing(p)
        p.paragraph_format.space_after = Pt(10)
        _style_run(p.add_run(summary), size_pt=8.5, color=C_TEXT_BODY)

    skills = [str(s).strip() for s in (data.get("skills") or []) if str(s).strip()]
    if skills:
        _add_section_title_row(doc, "Навыки", accent=M_ACCENT, dark="111827")
        _add_skills_chips(
            doc,
            skills,
            bg=M_CHIP_BG,
            border=M_CHIP_BORDER,
            text=M_CHIP_TEXT,
            font_pt=8,
            cols=3,
        )

    achievements = [str(a).strip() for a in (data.get("key_achievements") or []) if str(a).strip()]
    if achievements:
        _add_section_title_row(doc, "Ключевые достижения", accent=M_ACCENT, dark="111827")
        _add_achievement_entries(doc, achievements, accent=M_ACCENT)

    experience = data.get("experience") or []
    if experience:
        _add_section_title_row(doc, "Опыт работы", accent=M_ACCENT, dark="111827")
        _add_experience_entries(
            doc,
            experience,
            accent=M_ACCENT,
            position_color=M_ACCENT,
            company_color="111827",
            tab_cm=15,
            company_pt=9,
            position_pt=8,
            bullet_pt=8,
        )

    education = data.get("education") or []
    if education:
        _add_section_title_row(doc, "Образование", accent=M_ACCENT, dark="111827")
        _add_education_entries(doc, education, institution_pt=8.5, details_pt=7.5)

    langs = _non_native_languages(data)
    if langs:
        _add_section_title_row(doc, "Языки", accent=M_ACCENT, dark="111827")
        for lang in data.get("languages") or []:
            lang_s = str(lang).strip()
            if lang_s:
                p = doc.add_paragraph()
                _clear_paragraph_spacing(p)
                p.paragraph_format.space_after = Pt(2)
                _style_run(p.add_run(lang_s), size_pt=8, color=C_TEXT_BODY)

    docs = [str(d).strip() for d in (data.get("documents_and_permits") or []) if str(d).strip()]
    if docs:
        _add_section_title_row(doc, "Документы и допуски", accent=M_ACCENT, dark="111827")
        for doc_line in docs:
            p = doc.add_paragraph()
            _clear_paragraph_spacing(p)
            p.paragraph_format.space_after = Pt(2)
            _style_run(p.add_run(doc_line), size_pt=8, color=C_TEXT_BODY)

    return _save(doc)


def _fill_compact_sidebar(cell, data: dict[str, Any]) -> None:
    _set_cell_bg(cell, K_SIDEBAR_BG)
    _set_cell_margins(cell, top=160, start=140, bottom=160, end=140)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

    _add_light_sidebar_title(cell, "Контакты")
    for label, key in (("Телефон", "phone"), ("Email", "email"), ("Город", "city")):
        val = str(data.get(key) or "").strip()
        if val:
            _add_sidebar_label_value(cell, label, val, muted="9CA3AF", text="374151")

    salary = str(data.get("salary") or "").strip()
    if salary:
        _add_compact_salary_box(cell, salary)

    sched = _schedule_text(data)
    if sched:
        _add_light_sidebar_title(cell, "График")
        p = cell.add_paragraph()
        _clear_paragraph_spacing(p)
        p.paragraph_format.left_indent = Cm(0.25)
        p.paragraph_format.first_line_indent = Cm(-0.25)
        _style_run(p.add_run("· "), size_pt=8, bold=True, color=K_ACCENT)
        _style_run(p.add_run(sched), size_pt=7, color="374151")

    relocation = str(data.get("relocation") or "").strip()
    if relocation:
        _add_light_sidebar_title(cell, "Переезд")
        p = cell.add_paragraph()
        _clear_paragraph_spacing(p)
        p.paragraph_format.left_indent = Cm(0.25)
        p.paragraph_format.first_line_indent = Cm(-0.25)
        _style_run(p.add_run("· "), size_pt=8, bold=True, color=K_ACCENT)
        _style_run(p.add_run(relocation), size_pt=7, color="374151")

    skills = [str(s).strip() for s in (data.get("skills") or []) if str(s).strip()]
    if skills:
        _add_light_sidebar_title(cell, "Навыки")
        _add_skills_chips(
            cell,
            skills,
            bg="EDE9FE",
            border="DDD6FE",
            text="4C1D95",
            font_pt=6.5,
            cols=2,
        )

    langs = _non_native_languages(data)
    if langs:
        _add_light_sidebar_title(cell, "Языки")
        for lang in data.get("languages") or []:
            lang_s = str(lang).strip()
            if lang_s:
                p = cell.add_paragraph()
                _clear_paragraph_spacing(p)
                p.paragraph_format.left_indent = Cm(0.25)
                p.paragraph_format.first_line_indent = Cm(-0.25)
                _style_run(p.add_run("· "), size_pt=8, bold=True, color=K_ACCENT)
                _style_run(p.add_run(lang_s), size_pt=7, color="374151")

    docs = [str(d).strip() for d in (data.get("documents_and_permits") or []) if str(d).strip()]
    if docs:
        _add_light_sidebar_title(cell, "Документы и допуски")
        for doc_line in docs:
            p = cell.add_paragraph()
            _clear_paragraph_spacing(p)
            p.paragraph_format.left_indent = Cm(0.25)
            p.paragraph_format.first_line_indent = Cm(-0.25)
            _style_run(p.add_run("· "), size_pt=8, bold=True, color=K_ACCENT)
            _style_run(p.add_run(doc_line), size_pt=7, color="374151")


def _fill_compact_main(cell, data: dict[str, Any]) -> None:
    _set_cell_margins(cell, top=160, start=140, bottom=160, end=140)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

    name = str(data.get("full_name") or "").strip()
    position = str(data.get("target_position") or "").strip()
    if name:
        p = cell.add_paragraph()
        _clear_paragraph_spacing(p)
        p.paragraph_format.space_after = Pt(2)
        _style_run(p.add_run(name), size_pt=16, bold=True, color="111827")
    if position:
        p = cell.add_paragraph()
        _clear_paragraph_spacing(p)
        p.paragraph_format.space_after = Pt(6)
        _style_run(p.add_run(position), size_pt=9, bold=True, color=K_ACCENT)
    _para_divider(cell, color=K_ACCENT, size=12, space_after=12)

    summary = str(data.get("summary") or "").strip()
    if summary:
        _add_main_title(cell, "О себе", accent=K_ACCENT, dark="111827")
        p = cell.add_paragraph()
        _clear_paragraph_spacing(p)
        p.paragraph_format.space_after = Pt(8)
        _style_run(p.add_run(summary), size_pt=8, color=C_TEXT_BODY)

    experience = data.get("experience") or []
    if experience:
        _add_main_title(cell, "Опыт работы", accent=K_ACCENT, dark="111827")
        _add_experience_entries(
            cell,
            experience,
            accent=K_ACCENT,
            position_color=K_ACCENT,
            company_color="111827",
            tab_cm=10.2,
            company_pt=8.5,
            position_pt=7.5,
            bullet_pt=7.5,
        )

    achievements = [str(a).strip() for a in (data.get("key_achievements") or []) if str(a).strip()]
    if achievements:
        _add_main_title(cell, "Ключевые достижения", accent=K_ACCENT, dark="111827")
        _add_achievement_entries(cell, achievements, accent=K_ACCENT, size_pt=7.5)

    education = data.get("education") or []
    if education:
        _add_main_title(cell, "Образование", accent=K_ACCENT, dark="111827")
        _add_education_entries(cell, education, institution_pt=8, details_pt=7)


def _build_compact_docx(data: dict[str, Any]) -> bytes:
    doc = Document()
    _apply_default_styles(doc)
    _apply_doc_properties(doc, data)
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(0.8)
    section.right_margin = Cm(0.8)
    section.top_margin = Cm(0.8)
    section.bottom_margin = Cm(0.8)

    table = doc.add_table(rows=1, cols=2)
    _remove_table_borders(table)
    _set_table_fixed_width(table)
    _set_row_min_height(table.rows[0])
    sidebar, main = table.rows[0].cells
    sidebar.width = Cm(6.0)
    main.width = Cm(13.0)

    _fill_compact_sidebar(sidebar, data)
    _fill_compact_main(main, data)
    return _save(doc)


def _save(doc: Document) -> bytes:
    buffer = io.BytesIO()
    doc.save(buffer)
    return embed_nunito_fonts(buffer.getvalue())


def generate_docx_bytes(resume_data: dict[str, Any], template_name: str = "classic") -> bytes:
    """Build DOCX matching the selected PDF template layout (hh.ru-ready)."""
    from services.font_assets import ensure_fonts

    ensure_fonts()
    data = _normalize_salary(normalize_resume_data(resume_data))
    template = (template_name or "classic").strip().lower()
    if template not in VALID_TEMPLATES:
        template = "classic"
    if template == "modern":
        return _build_modern_docx_clean(data)
    if template == "compact":
        return _build_compact_docx(data)
    return _build_classic_docx(data)
