import io

import pytest
from docx import Document

from services.docx_service import generate_docx_bytes

SAMPLE = {
    "full_name": "Иван Петров",
    "target_position": "Backend Developer",
    "city": "Москва",
    "phone": "+7 900 000-00-00",
    "email": "ivan@example.com",
    "salary": "150000",
    "work_schedule": ["Полный день", "Удалённо"],
    "relocation": "Готов к переезду",
    "summary": "Опытный разработчик с фокусом на Python и API.",
    "skills": ["Python", "FastAPI", "PostgreSQL"],
    "key_achievements": ["Ускорил API в 2 раза", "Внедрил CI/CD"],
    "experience": [
        {
            "company": "Tech Corp",
            "position": "Senior Backend",
            "period": "2020 — н.в.",
            "description": "Проектировал API\nОптимизировал запросы",
        }
    ],
    "education": [
        {"institution": "МГУ", "degree": "Прикладная математика", "year": "2018"}
    ],
    "languages": ["Русский — родной", "English — B2"],
    "documents_and_permits": ["Права категории B"],
}


@pytest.mark.parametrize("template", ["classic", "modern", "compact"])
def test_generate_docx_bytes_smoke(template: str) -> None:
    raw = generate_docx_bytes(SAMPLE, template)
    assert isinstance(raw, bytes)
    assert len(raw) > 1000
    doc = Document(io.BytesIO(raw))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Иван Петров" in text or any("Иван" in c.text for t in doc.tables for r in t.rows for c in r.cells)


def test_classic_docx_has_two_column_table() -> None:
    raw = generate_docx_bytes(SAMPLE, "classic")
    doc = Document(io.BytesIO(raw))
    assert len(doc.tables) == 1
    assert len(doc.tables[0].columns) == 2
    sidebar_text = doc.tables[0].rows[0].cells[0].text
    main_text = doc.tables[0].rows[0].cells[1].text
    assert "КОНТАКТЫ" in sidebar_text
    assert "Python" in sidebar_text
    assert "ОПЫТ РАБОТЫ" in main_text
    assert "Tech Corp" in main_text
    assert "+7 900" in main_text


def test_docx_filename_hh_prefix() -> None:
    from services.docx_service import docx_filename

    assert docx_filename({"full_name": "Иван Петров"}).startswith("Rezyume_")
    assert docx_filename({"full_name": "Иван Петров"}).endswith("_hh.docx")


def test_embedded_fonts_increase_size() -> None:
    from pathlib import Path

    fonts_dir = Path(__file__).resolve().parent.parent / "fonts"
    if not (fonts_dir / "NunitoSans-Regular.ttf").exists():
        pytest.skip("Nunito fonts not present")
    raw = generate_docx_bytes(SAMPLE, "classic")
    assert raw[:2] == b"PK"
    assert len(raw) > 50_000
